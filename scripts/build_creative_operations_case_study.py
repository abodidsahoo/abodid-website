from __future__ import annotations

import math
import os
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BUILD_DIR = ROOT / "artifacts" / "creative-operations-case-study"
ASSET_DIR = BUILD_DIR / "assets"
OUTPUT_PATH = BUILD_DIR / "creative-operations-platform-case-study.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

INK = "202124"
MUTED = "5F6368"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8F0FE"
PALE = "F7F9FC"
LINE = "DADCE0"
GREEN = "137333"
AMBER = "A15C00"
PURPLE = "7E57C2"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.lstrip("#")
    return RGBColor(int(value[:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def find_font(bold: bool = False) -> str:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if os.path.exists(candidate):
            return candidate
    return candidates[-1]


FONT_REGULAR = find_font(False)
FONT_BOLD = find_font(True)


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    *,
    fill: str,
    outline: str = LINE,
    width: int = 3,
    radius: int = 28,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=width)


def centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    *,
    size: int,
    fill: str = INK,
    bold: bool = False,
) -> None:
    font = pil_font(size, bold)
    left, top, right, bottom = box
    text_box = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=8)
    width = text_box[2] - text_box[0]
    height = text_box[3] - text_box[1]
    draw.multiline_text(
        ((left + right - width) / 2, (top + bottom - height) / 2),
        text,
        font=font,
        fill=f"#{fill}",
        align="center",
        spacing=8,
    )


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    *,
    fill: str = BLUE,
    width: int = 5,
) -> None:
    draw.line([start, end], fill=f"#{fill}", width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    left = (
        end[0] - length * math.cos(angle - math.pi / 6),
        end[1] - length * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - length * math.cos(angle + math.pi / 6),
        end[1] - length * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=f"#{fill}")


def make_placeholder(
    filename: str,
    title: str,
    guidance: str,
    *,
    width: int = 1600,
    height: int = 620,
) -> Path:
    path = ASSET_DIR / filename
    image = Image.new("RGB", (width, height), f"#{PALE}")
    draw = ImageDraw.Draw(image)

    inset = 30
    dash = 28
    gap = 18
    x = inset
    while x < width - inset:
        draw.line((x, inset, min(x + dash, width - inset), inset), fill=f"#{BLUE}", width=3)
        draw.line((x, height - inset, min(x + dash, width - inset), height - inset), fill=f"#{BLUE}", width=3)
        x += dash + gap
    y = inset
    while y < height - inset:
        draw.line((inset, y, inset, min(y + dash, height - inset)), fill=f"#{BLUE}", width=3)
        draw.line((width - inset, y, width - inset, min(y + dash, height - inset)), fill=f"#{BLUE}", width=3)
        y += dash + gap

    draw.rounded_rectangle(
        (width // 2 - 46, height // 2 - 125, width // 2 + 46, height // 2 - 33),
        radius=20,
        fill=f"#{LIGHT_BLUE}",
    )
    draw.line(
        (width // 2, height // 2 - 104, width // 2, height // 2 - 54),
        fill=f"#{BLUE}",
        width=6,
    )
    draw.line(
        (width // 2 - 25, height // 2 - 79, width // 2 + 25, height // 2 - 79),
        fill=f"#{BLUE}",
        width=6,
    )

    title_font = pil_font(44, True)
    guidance_font = pil_font(25, False)
    title_box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(
        ((width - (title_box[2] - title_box[0])) / 2, height // 2 - 5),
        title,
        font=title_font,
        fill=f"#{INK}",
    )
    guidance_box = draw.multiline_textbbox(
        (0, 0),
        guidance,
        font=guidance_font,
        align="center",
        spacing=7,
    )
    draw.multiline_text(
        (
            (width - (guidance_box[2] - guidance_box[0])) / 2,
            height // 2 + 65,
        ),
        guidance,
        font=guidance_font,
        fill=f"#{MUTED}",
        align="center",
        spacing=7,
    )
    image.save(path)
    return path


def make_mind_map() -> Path:
    path = ASSET_DIR / "idea-map.png"
    image = Image.new("RGB", (1600, 760), f"#{WHITE}")
    draw = ImageDraw.Draw(image)

    center = (800, 370)
    center_radius = 112
    nodes = [
        ((240, 150), "Publish", "Projects · writing"),
        ((800, 105), "Manage", "Content · media"),
        ((1360, 150), "Distribute", "Email · resources"),
        ((1360, 575), "Understand", "Analytics · SEO"),
        ((800, 650), "Research", "Notes · retrieval"),
        ((240, 575), "Experiment", "AI · gesture"),
    ]

    for position, _, _ in nodes:
        end = (position[0], position[1])
        vector = (end[0] - center[0], end[1] - center[1])
        length = math.hypot(*vector)
        start = (
            int(center[0] + vector[0] / length * center_radius),
            int(center[1] + vector[1] / length * center_radius),
        )
        finish = (
            int(end[0] - vector[0] / length * 120),
            int(end[1] - vector[1] / length * 62),
        )
        draw.line((start, finish), fill=f"#{LINE}", width=5)

    draw.ellipse(
        (
            center[0] - center_radius,
            center[1] - center_radius,
            center[0] + center_radius,
            center[1] + center_radius,
        ),
        fill=f"#{BLUE}",
    )
    centered_text(
        draw,
        (
            center[0] - center_radius,
            center[1] - center_radius,
            center[0] + center_radius,
            center[1] + center_radius,
        ),
        "Creative\nwork",
        size=42,
        fill=WHITE,
        bold=True,
    )

    for position, title, subtitle in nodes:
        x, y = position
        box = (x - 145, y - 64, x + 145, y + 64)
        rounded_box(draw, box, fill=PALE, outline=LINE, width=3, radius=24)
        centered_text(draw, (box[0], box[1] + 4, box[2], box[1] + 64), title, size=31, bold=True)
        centered_text(
            draw,
            (box[0], box[1] + 53, box[2], box[3] - 4),
            subtitle,
            size=21,
            fill=MUTED,
        )

    draw.text(
        (55, 700),
        "External services: Supabase · Cloudflare R2 · Resend · OpenRouter · Vercel",
        font=pil_font(22, False),
        fill=f"#{MUTED}",
    )
    image.save(path)
    return path


def make_lifecycle_flow() -> Path:
    path = ASSET_DIR / "publishing-flow.png"
    image = Image.new("RGB", (1600, 340), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    labels = [
        ("Create", "Structured blocks"),
        ("Preview", "Responsive states"),
        ("Validate", "Rules + conflicts"),
        ("Publish", "Immutable snapshot"),
        ("Deliver", "Public views + CDN"),
        ("Learn", "Analytics + feedback"),
    ]
    node_width = 210
    gap = 44
    start_x = 42
    top = 82
    for index, (title, subtitle) in enumerate(labels):
        left = start_x + index * (node_width + gap)
        box = (left, top, left + node_width, top + 150)
        fill = LIGHT_BLUE if index in (0, 3, 5) else PALE
        rounded_box(draw, box, fill=fill, outline=BLUE if index in (0, 3, 5) else LINE, radius=24)
        centered_text(draw, (left, top + 20, left + node_width, top + 76), title, size=29, bold=True)
        centered_text(
            draw,
            (left + 10, top + 72, left + node_width - 10, top + 132),
            subtitle,
            size=20,
            fill=MUTED,
        )
        if index < len(labels) - 1:
            draw_arrow(
                draw,
                (left + node_width + 8, top + 75),
                (left + node_width + gap - 8, top + 75),
                fill=BLUE,
                width=4,
            )
    image.save(path)
    return path


def make_architecture_diagram() -> Path:
    path = ASSET_DIR / "system-architecture.png"
    image = Image.new("RGB", (1600, 720), f"#{WHITE}")
    draw = ImageDraw.Draw(image)

    layer_titles = [
        ("People", ["Creator / operator", "Collaborator", "Audience / client"]),
        ("Experience", ["Astro public site", "React admin apps", "Research labs"]),
        ("Workflows", ["Publish", "Communicate", "Measure", "Retrieve"]),
        (
            "Cloud primitives",
            ["Supabase: data + auth", "R2 + Worker: media", "Resend + AI APIs", "Vercel: delivery"],
        ),
    ]
    colors = [PALE, LIGHT_BLUE, "EEF7EE", "FFF6E5"]
    outlines = [LINE, BLUE, GREEN, AMBER]
    x_positions = [35, 415, 795, 1175]
    box_width = 335
    top = 95
    bottom = 650

    for index, ((title, items), fill, outline, left) in enumerate(
        zip(layer_titles, colors, outlines, x_positions)
    ):
        box = (left, top, left + box_width, bottom)
        rounded_box(draw, box, fill=fill, outline=outline, width=4, radius=28)
        draw.text(
            (left + 26, top + 28),
            title,
            font=pil_font(33, True),
            fill=f"#{INK}",
        )
        y = top + 108
        for item in items:
            draw.rounded_rectangle(
                (left + 24, y, left + box_width - 24, y + 78),
                radius=17,
                fill=f"#{WHITE}",
                outline=f"#{LINE}",
                width=2,
            )
            centered_text(
                draw,
                (left + 34, y + 6, left + box_width - 34, y + 72),
                item,
                size=22,
                fill=INK,
                bold=False,
            )
            y += 100
        if index < len(layer_titles) - 1:
            draw_arrow(
                draw,
                (left + box_width + 8, (top + bottom) // 2),
                (x_positions[index + 1] - 10, (top + bottom) // 2),
                fill=BLUE,
                width=5,
            )

    draw.text(
        (40, 25),
        "A modular platform: people see tasks, while the system coordinates services.",
        font=pil_font(25, False),
        fill=f"#{MUTED}",
    )
    image.save(path)
    return path


def make_module_grid() -> Path:
    path = ASSET_DIR / "module-grid.png"
    image = Image.new("RGB", (1600, 610), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    modules = [
        ("Publishing", "Portfolio CMS\nDrafts · preview · versions", BLUE),
        ("Operations", "Admin workspace\nMedia · accounts · XR", DARK_BLUE),
        ("Communication", "Newsletter + contact\nResend · subscriber state", GREEN),
        ("Insights", "Analytics + SEO\nJourneys · metadata · OG", AMBER),
        ("Knowledge", "Obsidian RAG\nHybrid retrieval · citations", PURPLE),
        ("Experiments", "Gesture + vision\nMediaPipe · TensorFlow.js", "C23B22"),
    ]
    margin_x = 35
    gap_x = 30
    gap_y = 28
    box_w = (1600 - margin_x * 2 - gap_x * 2) // 3
    box_h = 250
    for index, (title, detail, accent) in enumerate(modules):
        row = index // 3
        col = index % 3
        left = margin_x + col * (box_w + gap_x)
        top = 25 + row * (box_h + gap_y)
        box = (left, top, left + box_w, top + box_h)
        rounded_box(draw, box, fill=PALE, outline=LINE, width=3, radius=25)
        draw.rounded_rectangle(
            (left + 24, top + 26, left + 76, top + 78),
            radius=15,
            fill=f"#{accent}",
        )
        draw.text(
            (left + 98, top + 30),
            title,
            font=pil_font(30, True),
            fill=f"#{INK}",
        )
        draw.multiline_text(
            (left + 28, top + 112),
            detail,
            font=pil_font(22, False),
            fill=f"#{MUTED}",
            spacing=12,
        )
    image.save(path)
    return path


def make_artifact_placeholder() -> Path:
    path = ASSET_DIR / "design-artifacts-placeholder.png"
    image = Image.new("RGB", (1600, 590), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    labels = [
        ("01", "Storyboard", "Show the user’s context\nand moment of need"),
        ("02", "Wireframe", "Show task hierarchy\nand system states"),
        ("03", "Prototype", "Show interaction,\nfeedback and recovery"),
    ]
    margin = 35
    gap = 32
    box_w = (1600 - margin * 2 - gap * 2) // 3
    for index, (number, title, guidance) in enumerate(labels):
        left = margin + index * (box_w + gap)
        box = (left, 30, left + box_w, 550)
        rounded_box(draw, box, fill=PALE, outline=LINE, width=3, radius=24)
        draw.text((left + 28, 58), number, font=pil_font(23, True), fill=f"#{BLUE}")
        draw.text((left + 28, 101), title, font=pil_font(32, True), fill=f"#{INK}")
        for y in (190, 272, 354):
            draw.rounded_rectangle(
                (left + 28, y, left + box_w - 28, y + 55),
                radius=12,
                fill=f"#{WHITE}",
                outline=f"#{LINE}",
                width=2,
            )
        draw.multiline_text(
            (left + 28, 442),
            guidance,
            font=pil_font(20, False),
            fill=f"#{MUTED}",
            spacing=8,
        )
    image.save(path)
    return path


def make_roadmap() -> Path:
    path = ASSET_DIR / "enterprise-roadmap.png"
    image = Image.new("RGB", (1600, 360), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    steps = [
        ("1", "Working prototype", "Single operator\nreal workflows"),
        ("2", "Team validation", "Roles · audit trail\nusability studies"),
        ("3", "Enterprise platform", "Multi-tenancy · SSO\nSLOs · governance"),
    ]
    lefts = [75, 570, 1065]
    for index, ((number, title, detail), left) in enumerate(zip(steps, lefts)):
        box = (left, 62, left + 405, 292)
        rounded_box(
            draw,
            box,
            fill=LIGHT_BLUE if index == 0 else PALE,
            outline=BLUE if index == 0 else LINE,
            width=4 if index == 0 else 3,
            radius=25,
        )
        draw.ellipse((left + 25, 88, left + 87, 150), fill=f"#{BLUE}")
        centered_text(draw, (left + 25, 88, left + 87, 150), number, size=26, fill=WHITE, bold=True)
        draw.text((left + 108, 88), title, font=pil_font(29, True), fill=f"#{INK}")
        draw.multiline_text(
            (left + 108, 152),
            detail,
            font=pil_font(21, False),
            fill=f"#{MUTED}",
            spacing=9,
        )
        if index < 2:
            draw_arrow(draw, (left + 420, 177), (lefts[index + 1] - 18, 177), fill=BLUE, width=5)
    image.save(path)
    return path


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str = INK,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(
    paragraph,
    *,
    before: float = 0,
    after: float = 6,
    line: float = 1.25,
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def paragraph_border(
    paragraph,
    *,
    side: str = "left",
    color: str = BLUE,
    size: int = 18,
    space: int = 8,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def keep_with_next(paragraph, value: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    element = p_pr.find(qn("w:keepNext"))
    if value and element is None:
        element = OxmlElement("w:keepNext")
        p_pr.append(element)


def configure_document(doc: Document) -> tuple[int, int]:
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    apply_section_geometry(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    bullet_num_id = add_numbering_definition(doc, bullet=True)
    decimal_num_id = add_numbering_definition(doc, bullet=False)
    configure_header_footer(section)
    return bullet_num_id, decimal_num_id


def apply_section_geometry(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False


def add_numbering_definition(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
        if element.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
        if element.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    if bullet:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Symbol")
        r_fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(r_fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    page_run = OxmlElement("w:r")
    page_text = OxmlElement("w:t")
    page_text.text = "1"
    page_run.append(page_text)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_char_begin)
    paragraph._p.append(instr_text)
    paragraph._p.append(fld_char_separate)
    paragraph._p.append(page_run)
    paragraph._p.append(fld_char_end)


def configure_header_footer(section) -> None:
    for header in (section.header, section.even_page_header, section.first_page_header):
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.text = ""
        set_paragraph_spacing(paragraph, after=0, line=1.0)

    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.text = ""
        set_paragraph_spacing(paragraph, after=0, line=1.0)


def add_kicker(doc: Document, text: str, *, after: float = 5) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=after, line=1.0)
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=9, color=BLUE, bold=True)
    keep_with_next(paragraph)


def add_title(doc: Document, text: str, subtitle: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=7, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=30, color=INK, bold=True)
    keep_with_next(paragraph)
    if subtitle:
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph, before=0, after=14, line=1.15)
        run = paragraph.add_run(subtitle)
        set_run_font(run, size=14, color=MUTED)
        keep_with_next(paragraph)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def add_body(
    doc: Document,
    text: str,
    *,
    bold_lead: str | None = None,
    after: float = 6,
    size: float = 10.3,
    color: str = INK,
) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=after, line=1.22)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=size, color=color, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, size=size, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=color)


def add_label_detail(
    doc: Document,
    label: str,
    detail: str,
    *,
    after: float = 3,
) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=after, line=1.14)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=9.5, color=BLUE, bold=True)
    detail_run = paragraph.add_run(detail)
    set_run_font(detail_run, size=9.5, color=INK)


def add_callout(doc: Document, text: str, *, label: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=4, after=10, line=1.2)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph_shading(paragraph, LIGHT_BLUE)
    paragraph_border(paragraph, side="left", color=BLUE, size=20, space=10)
    if label:
        run = paragraph.add_run(f"{label.upper()}\n")
        set_run_font(run, size=8.5, color=BLUE, bold=True)
    run = paragraph.add_run(text)
    set_run_font(run, size=11.2, color=INK, bold=True)


def add_bullet(doc: Document, num_id: int, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    set_paragraph_spacing(paragraph, after=4, line=1.25)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=9.6, color=INK, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, size=9.6, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=9.6, color=INK)


def add_picture(
    doc: Document,
    path: Path,
    *,
    width: float,
    caption: str | None = None,
    after: float = 6,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=2, after=2, line=1.0)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    if caption:
        caption_paragraph = doc.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(caption_paragraph, after=after, line=1.0)
        caption_run = caption_paragraph.add_run(caption)
        set_run_font(caption_run, size=8.2, color=MUTED, italic=True)


def add_code_block(doc: Document, lines: list[str], caption: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=3, after=3, line=1.0)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph_shading(paragraph, "F3F5F7")
    paragraph_border(paragraph, side="left", color=BLUE, size=16, space=8)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, name="Consolas", size=8.1, color=INK)
    caption_paragraph = doc.add_paragraph()
    set_paragraph_spacing(caption_paragraph, after=8, line=1.05)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=8.3, color=MUTED, italic=True)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths)):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 8.6) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=0, line=1.12)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_readiness_table(doc: Document) -> None:
    rows = [
        (
            "Access",
            "Supabase Auth, RLS and server-only privileged keys",
            "Organization tenancy, SSO and audit logs",
        ),
        (
            "Reliability",
            "Immutable publish backups, queue retries and recovery",
            "SLOs, incident runbooks and disaster recovery",
        ),
        (
            "Data",
            "Public/private views and signed media uploads",
            "Retention, governance and regional policy",
        ),
        (
            "Quality",
            "Focused unit and end-to-end tests",
            "CI gates, load tests and formal accessibility audit",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1700, 4000, 3660])
    headers = ["Capability", "Working prototype", "Enterprise validation"]
    for cell, text in zip(table.rows[0].cells, headers):
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, text, bold=True, color=DARK_BLUE, size=8.7)
    for capability, evidence, next_step in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], capability, bold=True, color=INK)
        set_cell_text(cells[1], evidence)
        set_cell_text(cells[2], next_step)
    set_table_geometry(table, [1700, 4000, 3660])
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=4, after=7, line=1.0)
    run = paragraph.add_run(
        "Positioning: enterprise-pattern product prototype—not a claim of enterprise deployment or scale."
    )
    set_run_font(run, size=8.2, color=MUTED, italic=True)


def add_page_break(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_section_geometry(section)
    configure_header_footer(section)


def build_document() -> Path:
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    hero_placeholder = make_placeholder(
        "hero-placeholder.png",
        "SCREENSHOT 01 — PRODUCT OVERVIEW",
        "Use the admin overview or a composed view of the public site + control room.",
        height=560,
    )
    research_placeholder = make_placeholder(
        "research-placeholder.png",
        "DESIGN EVIDENCE — NOTES + EARLY FLOWS",
        "Add inquiry themes, a handwritten system map, and first wireframes.",
        height=390,
    )
    editor_placeholder = make_placeholder(
        "editor-placeholder.png",
        "SCREENSHOT 02 — WORK EDITOR",
        "Show the block list, live preview, draft state, and publish controls.",
        height=500,
    )
    app_placeholder = make_placeholder(
        "apps-placeholder.png",
        "SCREENSHOTS 03–06 — KEY APPLICATIONS",
        "Newsletter studio · analytics · media library · Obsidian assistant",
        height=400,
    )
    final_placeholder = make_placeholder(
        "final-placeholder.png",
        "FINAL PRODUCT MOMENT",
        "Add one strong high-fidelity screen that demonstrates clarity at system scale.",
        height=370,
    )

    mind_map = make_mind_map()
    lifecycle = make_lifecycle_flow()
    architecture = make_architecture_diagram()
    module_grid = make_module_grid()
    artifact_placeholder = make_artifact_placeholder()
    roadmap = make_roadmap()

    doc = Document()
    bullet_num_id, _ = configure_document(doc)

    # Page 1 — Cover and product position
    add_kicker(doc, "Interaction design case study")
    add_title(
        doc,
        "Creative Operations Platform",
        "From a personal portfolio to a unified, enterprise-pattern product prototype.",
    )
    add_callout(
        doc,
        "One workspace to publish, distribute, understand, and evolve creative work.",
        label="Product promise",
    )
    add_picture(
        doc,
        hero_placeholder,
        width=6.25,
        caption="Hero image placeholder · Recommended crop: 16:7",
        after=8,
    )
    add_label_detail(doc, "ROLE", "Product strategy · Interaction design · Visual design · Full-stack prototyping")
    add_label_detail(doc, "STAGE", "Working single-operator platform · Enterprise validation roadmap defined")
    add_label_detail(doc, "CORE STACK", "Astro · React · TypeScript · Supabase · Cloudflare R2 · Vercel · Resend")
    add_label_detail(doc, "SCOPE", "Public site · Admin workspace · AI tools · Media pipeline · Analytics")
    add_body(
        doc,
        "The project began with a personal need: manage a growing body of work and repeated enquiries without switching between disconnected tools. The same workflow problem appears across independent creators, studios, and small teams.",
        size=9.6,
        color=MUTED,
        after=0,
    )

    # Page 2 — Problem, user signal and evolution
    add_page_break(doc)
    add_kicker(doc, "01 · Opportunity")
    add_heading(doc, "One creative practice. Too many disconnected tools.", 1)
    add_body(
        doc,
        "Publishing, files, newsletters, analytics, research notes, and enquiries lived in separate systems. Every hand-off lost context. Repeated inbound questions suggested that this was not only a personal workflow issue.",
        size=10.2,
    )
    add_label_detail(
        doc,
        "USER SIGNAL",
        "Repeated enquiries about projects, services, workshops, resources, and collaboration.",
        after=4,
    )
    add_label_detail(
        doc,
        "CORE NEED",
        "A single place to create, manage, publish, communicate, and learn.",
        after=4,
    )
    add_label_detail(
        doc,
        "PRODUCT HYPOTHESIS",
        "A modular operations platform can reduce context switching while preserving creative flexibility.",
        after=8,
    )
    add_picture(
        doc,
        mind_map,
        width=5.9,
        caption="The product grew outward from one central object: the work itself.",
        after=6,
    )
    add_picture(
        doc,
        research_placeholder,
        width=5.9,
        caption="Replace this placeholder with evidence, not decoration.",
        after=2,
    )
    add_body(
        doc,
        "Research boundary: current evidence combines direct workflow observation and recurring enquiry patterns. The next phase requires structured interviews, competitor analysis, and task-based usability testing.",
        size=8.8,
        color=MUTED,
        after=0,
    )

    # Page 3 — Core workflow and implementation
    add_page_break(doc)
    add_kicker(doc, "02 · Core product flow")
    add_heading(doc, "Design the entire content lifecycle.", 1)
    add_body(
        doc,
        "The key design decision was to treat publishing as a stateful workflow—not a save button. Every project moves through creation, preview, validation, publication, delivery, and learning.",
        size=10.2,
    )
    add_picture(
        doc,
        lifecycle,
        width=6.25,
        caption="Primary workflow · Each step exposes clear status, feedback, and recovery.",
        after=5,
    )
    add_picture(
        doc,
        editor_placeholder,
        width=6.25,
        caption="Show one task, one decision, and one visible system state.",
        after=5,
    )
    add_code_block(
        doc,
        [
            "await requirePortfolioAdmin();",
            'const { data, error } = await supabase.rpc("portfolio_save_draft", {',
            "  p_project_id: projectId,",
            "  p_expected_lock_version: draft.lockVersion,",
            "  p_payload: toSavePayload(draft),",
            "});",
        ],
        "The editor delegates authorization and conflict protection to a database function. A stale tab cannot overwrite newer work.",
    )
    add_bullet(
        doc,
        bullet_num_id,
        "Recovery by design. Published versions are immutable backups; restoring one never changes the live project until it is published again.",
        bold_lead="Recovery by design.",
    )
    add_bullet(
        doc,
        bullet_num_id,
        "Safe public surface. Sanitized views expose only published fields; draft and admin data remain private.",
        bold_lead="Safe public surface.",
    )

    # Page 4 — Architecture and applications
    add_page_break(doc)
    add_kicker(doc, "03 · Platform architecture")
    add_heading(doc, "Show tasks. Hide infrastructure complexity.", 1)
    add_body(
        doc,
        "Users interact with clear workflows. The platform coordinates data, identity, object storage, background processing, email, AI models, and delivery behind those workflows.",
        size=10.2,
    )
    add_picture(
        doc,
        architecture,
        width=6.25,
        caption="System model · Experience, workflow, and cloud layers remain modular.",
        after=4,
    )
    add_picture(
        doc,
        module_grid,
        width=6.25,
        caption="Six product modules share identity, media, content, and system feedback.",
        after=4,
    )
    add_code_block(
        doc,
        [
            "const questionEmbedding = await createEmbedding(searchQuery);",
            'const { data } = await supabase.rpc("match_obsidian_chunks", {',
            "  query_embedding: questionEmbedding,",
            "  match_count: 12,",
            "  match_threshold: 0.16,",
            "  public_only: publicOnly,",
            "});",
        ],
        "The knowledge assistant retrieves only permitted note chunks before it generates a cited answer.",
    )
    # Page 5 — Interaction principles and enterprise readiness
    add_page_break(doc)
    add_kicker(doc, "04 · Interaction design")
    add_heading(doc, "Make complex systems feel calm.", 1)
    add_body(
        doc,
        "The interface is organized around tasks and system states. Each screen should answer three questions: What can I do? What is happening? How can I recover?",
        size=10.2,
    )
    add_picture(
        doc,
        artifact_placeholder,
        width=6.25,
        caption="Portfolio artifact placeholders · Replace with real storyboard, wireframe, and prototype frames.",
        after=5,
    )
    principles = [
        ("Progressive disclosure.", "Show essential controls first; reveal advanced settings when needed."),
        ("Consistent status language.", "Draft, processing, published, failed, and recovered mean the same across tools."),
        ("Reversible actions.", "Version history, protected media, and explicit publishing reduce destructive mistakes."),
        ("Accessible by default.", "Semantic controls, keyboard access, alt text, reduced motion, and responsive layouts."),
        ("Shared patterns.", "Layouts, blocks, tokens, validation, feedback, and empty states form the product language."),
    ]
    for lead, detail in principles:
        add_bullet(doc, bullet_num_id, f"{lead} {detail}", bold_lead=lead)
    add_heading(doc, "Enterprise patterns, clearly bounded", 2)
    add_readiness_table(doc)

    # Page 6 — Reflection, next steps and role alignment
    add_page_break(doc)
    add_kicker(doc, "05 · Reflection")
    add_heading(doc, "A working product—and a clearer design practice.", 1)
    add_body(
        doc,
        "The prototype proves that publishing, communication, media, analytics, knowledge, and experimentation can share one coherent system. It also exposed where exploration must become product discipline.",
        size=10.2,
    )
    lessons = [
        (
            "I over-normalized the first portfolio schema.",
            "I learned that a portable document snapshot can be safer and easier to author than maximum relational purity.",
        ),
        (
            "I let experiments grow into large components.",
            "I learned to extract state machines, hooks, primitives, and design tokens once an idea proves useful.",
        ),
        (
            "Quality gates arrived too late.",
            "I learned that tests, type checks, linting, builds, security checks, and documentation must run as one system.",
        ),
        (
            "Publishing exposed operational gaps.",
            "I learned that cache freshness, staging access, observability, and recovery are part of the user experience.",
        ),
    ]
    for mistake, lesson in lessons:
        add_label_detail(doc, "MISTAKE", mistake, after=1)
        add_label_detail(doc, "LESSON", lesson, after=5)
    add_picture(
        doc,
        roadmap,
        width=6.25,
        caption="Scale is a validation path, not a visual claim.",
        after=5,
    )
    add_heading(doc, "Why this aligns with Google ACI UX", 2)
    alignment_points = [
        "Translates complex technical systems into understandable user flows.",
        "Connects storyboards, wireframes, prototypes, specifications, and implementation.",
        "Balances user needs with security, reliability, storage, compute, and operational constraints.",
        "Demonstrates end-to-end ownership: concept, system design, build, measurement, and iteration.",
    ]
    for item in alignment_points:
        add_bullet(doc, bullet_num_id, item)
    add_callout(
        doc,
        "The domain is creative operations. The transferable skill is designing clarity, control, and trust across a complex cloud-backed system.",
        label="Role fit",
    )
    add_picture(
        doc,
        final_placeholder,
        width=6.25,
        caption="Closing visual placeholder · Add product link or QR code when the case study is published.",
        after=2,
    )

    doc.core_properties.title = "Creative Operations Platform — Interaction Design Case Study"
    doc.core_properties.subject = "Enterprise-pattern product prototype aligned to interaction design roles"
    doc.core_properties.author = "Abodid Sahoo"
    doc.core_properties.keywords = "UX, interaction design, enterprise prototype, creative technology"
    doc.core_properties.comments = ""
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_document()
    print(output)
